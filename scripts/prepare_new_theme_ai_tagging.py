import argparse
import csv
import hashlib
import json
import re
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
WORKSPACE = ROOT.parent
DEFAULT_DOCX = WORKSPACE / "themes new.docx"
DEFAULT_OUT_DIR = WORKSPACE / "theme_tagging_new"
CRS_RECORDS_INDEX = ROOT / "public" / "data" / "crs-decade-records" / "index.json"

TEXT_FIELDS = [
    "title",
    "short_description",
    "long_description",
    "purpose",
    "mode",
    "mode_detail",
    "flow",
]

METADATA_FIELDS = [
    "row_number",
    "year",
    "donor",
    "agency",
    "recipient",
    "recipient_scope",
    "recipient_region_detail",
    "region",
    "mode",
    "mode_detail",
    "flow",
    "commitment_defl",
    "disbursement_defl",
    "climate_mitigation",
    "climate_adaptation",
    "gender",
    "drr",
    "biodiversity",
    "environment",
]

GENERIC_TEXT = {
    "",
    "na",
    "n/a",
    "none",
    "unknown",
    "transport",
    "road transport",
    "rail transport",
    "water transport",
    "air transport",
    "transport policy and administrative management",
    "education and training in transport and storage",
    "transport training and education",
    "transport and storage, unspecified",
    "storage and logistics",
    "ii.1. transport & storage",
    "other",
}

RECORDS_PER_TASK = 20


def clean(value):
    return re.sub(r"\s+", " ", str(value or "")).strip()


def slug(value):
    text = re.sub(r"[^a-z0-9]+", "_", value.lower()).strip("_")
    return text or "unknown"


def read_docx_definitions(path):
    doc = Document(path)
    paragraphs = [clean(paragraph.text) for paragraph in doc.paragraphs if clean(paragraph.text)]
    tables = doc.tables
    if len(tables) < 3:
        raise ValueError(f"Expected at least 3 definition tables in {path}, found {len(tables)}")

    active_definition = ""
    urban_definition = ""
    for index, paragraph in enumerate(paragraphs):
        lower = paragraph.lower()
        if lower == "active transport" and index + 1 < len(paragraphs):
            active_definition = paragraphs[index + 1]
        if lower == "urban transport" and index + 1 < len(paragraphs):
            urban_definition = paragraphs[index + 1]

    def parse_table(table):
        rows = []
        for row in table.rows[1:]:
            cells = [clean(cell.text) for cell in row.cells]
            if len(cells) < 3 or not cells[0]:
                continue
            keywords = [
                clean(keyword)
                for keyword in re.split(r"[;,]", cells[2])
                if clean(keyword)
            ]
            rows.append(
                {
                    "id": slug(cells[0]),
                    "label": cells[0],
                    "description": cells[1],
                    "keywords": keywords,
                }
            )
        return rows

    return [
        {
            "id": "e_mobility",
            "label": "E-mobility",
            "definition": (
                "Projects that explicitly support electric mobility, electric vehicles, charging or battery infrastructure, "
                "grid integration for e-mobility, electric freight and logistics, EV finance or market development, pilots, "
                "skills, data, or inclusion related to electric mobility."
            ),
            "subthemes": parse_table(tables[0]),
        },
        {
            "id": "active_transport",
            "label": "Active Transport",
            "definition": active_definition,
            "subthemes": parse_table(tables[1]),
        },
        {
            "id": "urban_transport",
            "label": "Urban Transport",
            "definition": urban_definition,
            "subthemes": parse_table(tables[2]),
        },
    ]


def definition_keywords(definitions):
    out = []
    for theme in definitions:
        out.append((theme["label"], theme["id"], "theme_label"))
        for subtheme in theme["subthemes"]:
            out.append((subtheme["label"], theme["id"], subtheme["id"]))
            for keyword in subtheme.get("keywords", []):
                out.append((keyword, theme["id"], subtheme["id"]))
    unique = []
    seen = set()
    for keyword, theme_id, subtheme_id in out:
        key = (keyword.lower(), theme_id, subtheme_id)
        if key not in seen:
            seen.add(key)
            unique.append((keyword, theme_id, subtheme_id))
    return unique


def keyword_in_text(keyword, text):
    keyword = clean(keyword).lower()
    if not keyword:
        return False
    if len(keyword) <= 3 and re.fullmatch(r"[a-z0-9]+", keyword):
        return re.search(rf"\b{re.escape(keyword)}\b", text) is not None
    return re.search(rf"\b{re.escape(keyword)}\b", text) is not None if " " not in keyword else keyword in text


def keyword_hits(record_text, keywords):
    text = record_text.lower()
    hits = []
    for keyword, theme_id, subtheme_id in keywords:
        if keyword_in_text(keyword, text):
            hits.append({"keyword": keyword, "theme_id": theme_id, "subtheme_id": subtheme_id})
    return hits


def load_records():
    index = json.loads(CRS_RECORDS_INDEX.read_text(encoding="utf-8"))
    records = []
    for chunk in index["chunks"]:
        chunk_path = ROOT / "public" / chunk["file"]
        records.extend(json.loads(chunk_path.read_text(encoding="utf-8")))
    return sorted(records, key=lambda row: row.get("row_number") or 0)


def field_is_meaningful(value):
    text = clean(value)
    return text.lower() not in GENERIC_TEXT


def record_text_parts(record):
    parts = []
    seen_values = set()
    for field in TEXT_FIELDS:
        value = clean(record.get(field))
        if not field_is_meaningful(value):
            continue
        key = value.lower()
        if key in seen_values:
            continue
        seen_values.add(key)
        parts.append({"field": field, "value": value})
    return parts


def record_text(record):
    parts = record_text_parts(record)
    return "\n".join(f"{part['field']}: {part['value']}" for part in parts)


def alpha_count(text):
    return len(re.findall(r"[A-Za-z]", text))


def word_count(text):
    return len(re.findall(r"[A-Za-z][A-Za-z-]*", text))


def classify_prefilter(record, keywords, mode):
    text = record_text(record)
    hits = [] if mode == "semantic" else keyword_hits(text, keywords)
    letters = alpha_count(text)
    words = word_count(text)

    if hits:
        return True, "keyword_signal", hits, letters, words

    if not text:
        return False, "no_usable_title_or_description", hits, letters, words

    if mode == "keyword":
        return False, "no_theme_keyword_signal", hits, letters, words

    has_narrative = any(field_is_meaningful(record.get(field)) for field in ["title", "short_description", "long_description"])
    if not has_narrative:
        return False, "no_project_narrative_fields", hits, letters, words

    if letters < 45 or words < 6:
        return False, "sparse_or_generic_text", hits, letters, words

    return True, "substantive_text_for_ai_review", hits, letters, words


def record_payload(record):
    return {
        "metadata": {field: record.get(field, "") for field in METADATA_FIELDS},
        "text_fields": {field: clean(record.get(field)) for field in TEXT_FIELDS if field_is_meaningful(record.get(field))},
    }


def payload_hash(payload):
    relevant = {
        "text_fields": payload["text_fields"],
        "mode": payload["metadata"].get("mode", ""),
        "mode_detail": payload["metadata"].get("mode_detail", ""),
        "flow": payload["metadata"].get("flow", ""),
    }
    raw = json.dumps(relevant, ensure_ascii=False, sort_keys=True)
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()[:16]


def compact_definitions_for_prompt(definitions):
    lines = []
    for theme in definitions:
        subtheme_ids = ", ".join(subtheme["id"] for subtheme in theme["subthemes"])
        lines.append(f"{theme['label']} (theme_id={theme['id']}): {theme['definition']}")
        lines.append(f"Allowed subtheme_id values for {theme['id']} only: {subtheme_ids}")
        for subtheme in theme["subthemes"]:
            lines.append(f"- {subtheme['label']} (subtheme_id={subtheme['id']}): {subtheme['description']}")
    return "\n".join(lines)


def response_schema(definitions):
    theme_ids = [theme["id"] for theme in definitions]
    subtheme_ids = sorted({subtheme["id"] for theme in definitions for subtheme in theme["subthemes"]})
    return {
        "type": "object",
        "additionalProperties": False,
        "required": ["records"],
        "properties": {
            "records": {
                "type": "array",
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "required": ["task_record_id", "tags", "needs_human_review", "review_reason"],
                    "properties": {
                        "task_record_id": {"type": "string"},
                        "tags": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "additionalProperties": False,
                                "required": ["theme_id", "confidence", "reason", "evidence_fields", "subthemes"],
                                "properties": {
                                    "theme_id": {"type": "string", "enum": theme_ids},
                                    "confidence": {"type": "string", "enum": ["high", "medium", "low"]},
                                    "reason": {"type": "string"},
                                    "evidence_fields": {"type": "array", "items": {"type": "string"}},
                                    "subthemes": {
                                        "type": "array",
                                        "items": {
                                            "type": "object",
                                            "additionalProperties": False,
                                            "required": ["subtheme_id", "confidence", "reason", "evidence_fields"],
                                            "properties": {
                                                "subtheme_id": {"type": "string", "enum": subtheme_ids},
                                                "confidence": {"type": "string", "enum": ["high", "medium", "low"]},
                                                "reason": {"type": "string"},
                                                "evidence_fields": {"type": "array", "items": {"type": "string"}},
                                            },
                                        },
                                    },
                                },
                            },
                        },
                        "needs_human_review": {"type": "boolean"},
                        "review_reason": {"type": "string"},
                    },
                },
            },
        },
    }


def tagging_instructions(definitions):
    return (
        "You are tagging OECD CRS transport finance records for a high-stakes analytical dashboard.\n"
        "Use semantic judgment and project context, not keyword matching. Be strict and evidence-based.\n"
        "Tag a theme only when the provided title, description, purpose, mode, or other row details explicitly support it.\n"
        "Do not infer from country, donor, amount, broad transport mode, climate marker, or general development language alone.\n"
        "Do not tag Active Transport merely because a project mentions road safety unless pedestrians, cyclists, walking, cycling, NMT, complete streets, or vulnerable-road-user measures are explicit.\n"
        "Do not tag Urban Transport merely because a project is a road project; require city, municipal, metropolitan, urban public transport, urban roads/streets, traffic/congestion, urban terminal, TOD, or similar explicit evidence.\n"
        "Do not tag E-mobility unless electric mobility, EVs, charging, batteries, grid links for EVs, electric freight, or EV market/capacity elements are explicit.\n"
        "A record may receive multiple themes and multiple subthemes, but only when each is supported by evidence.\n"
        "Subthemes must stay under their parent theme. Never put an Active Transport or E-mobility subtheme under an Urban Transport tag, and never put an Urban Transport subtheme under another theme.\n"
        "If evidence supports a subtheme from another theme, create a separate positive theme tag for that parent theme. If the parent theme itself is not clearly supported, omit that subtheme.\n"
        "Return only positive tags. If no theme applies, return tags: [] for that record.\n"
        "For every positive theme and subtheme, provide a short reason and evidence_fields. Do not provide negative-theme explanations.\n"
        "Mark needs_human_review=true when evidence is ambiguous, generic, contradictory, or only weakly implied.\n\n"
        "Theme definitions and subthemes:\n"
        f"{compact_definitions_for_prompt(definitions)}"
    )


def write_json(path, data):
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def write_jsonl(path, rows):
    with path.open("w", encoding="utf-8") as handle:
        for row in rows:
            handle.write(json.dumps(row, ensure_ascii=False, separators=(",", ":")) + "\n")


def write_csv(path, rows, fieldnames):
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def main():
    parser = argparse.ArgumentParser(description="Prepare CRS records for strict AI tagging under the new theme definitions.")
    parser.add_argument("--docx", type=Path, default=DEFAULT_DOCX)
    parser.add_argument("--out-dir", type=Path, default=DEFAULT_OUT_DIR)
    parser.add_argument("--filter-mode", choices=["semantic", "balanced", "keyword"], default="semantic")
    parser.add_argument("--records-per-task", type=int, default=RECORDS_PER_TASK)
    args = parser.parse_args()

    args.out_dir.mkdir(parents=True, exist_ok=True)
    definitions = read_docx_definitions(args.docx)
    keywords = definition_keywords(definitions)
    records = load_records()

    unique_payloads = {}
    record_map = []
    excluded = []
    reason_counts = Counter()
    keyword_theme_counts = Counter()
    candidate_reason_counts = Counter()

    for record in records:
        include, reason, hits, letters, words = classify_prefilter(record, keywords, args.filter_mode)
        reason_counts[reason] += 1
        if not include:
            excluded.append(
                {
                    "row_number": record.get("row_number", ""),
                    "year": record.get("year", ""),
                    "recipient": record.get("recipient", ""),
                    "title": record.get("title", ""),
                    "short_description": record.get("short_description", ""),
                    "long_description": record.get("long_description", ""),
                    "purpose": record.get("purpose", ""),
                    "prefilter_reason": reason,
                    "alpha_count": letters,
                    "word_count": words,
                }
            )
            continue

        candidate_reason_counts[reason] += 1
        for hit in hits:
            keyword_theme_counts[hit["theme_id"]] += 1

        payload = record_payload(record)
        task_id = payload_hash(payload)
        if task_id not in unique_payloads:
            unique_payloads[task_id] = {"task_record_id": task_id, **payload}

        record_map.append(
            {
                "row_number": record.get("row_number", ""),
                "task_record_id": task_id,
                "prefilter_reason": reason,
            }
        )

    unique_rows = sorted(unique_payloads.values(), key=lambda item: item["task_record_id"])
    tasks = []
    instructions = tagging_instructions(definitions)
    schema = response_schema(definitions)
    for index in range(0, len(unique_rows), args.records_per_task):
        group = unique_rows[index : index + args.records_per_task]
        tasks.append(
            {
                "custom_id": f"new-theme-tagging-{index // args.records_per_task + 1:05d}",
                "instructions": instructions,
                "response_schema": schema,
                "records": group,
            }
        )

    write_json(args.out_dir / "theme_definitions_new.json", {"source": str(args.docx), "themes": definitions})
    write_json(args.out_dir / "ai_response_schema.json", schema)
    summary_payload = {
        "source_records": len(records),
        "filter_mode": args.filter_mode,
        "records_per_task": args.records_per_task,
        "candidate_record_rows": len(record_map),
        "excluded_record_rows": len(excluded),
        "unique_ai_payloads_after_deduplication": len(unique_rows),
        "ai_task_groups": len(tasks),
        "prefilter_reason_counts": dict(reason_counts),
        "candidate_reason_counts": dict(candidate_reason_counts),
        "rough_character_volume_for_unique_payloads": sum(len(json.dumps(row, ensure_ascii=False)) for row in unique_rows),
        "rough_input_token_estimate_for_payloads_only": round(sum(len(json.dumps(row, ensure_ascii=False)) for row in unique_rows) / 4),
    }
    if args.filter_mode != "semantic":
        summary_payload["keyword_theme_hit_counts"] = dict(keyword_theme_counts)
    write_json(args.out_dir / "prefilter_summary.json", summary_payload)
    write_jsonl(args.out_dir / "unique_ai_payloads.jsonl", unique_rows)
    write_jsonl(args.out_dir / "ai_tagging_tasks.jsonl", tasks)
    write_csv(args.out_dir / "record_to_task_map.csv", record_map, ["row_number", "task_record_id", "prefilter_reason"])
    write_csv(
        args.out_dir / "excluded_records_prefilter.csv",
        excluded,
        ["row_number", "year", "recipient", "title", "short_description", "long_description", "purpose", "prefilter_reason", "alpha_count", "word_count"],
    )

    print(f"Definitions: {args.out_dir / 'theme_definitions_new.json'}")
    print(f"AI tasks: {args.out_dir / 'ai_tagging_tasks.jsonl'}")
    print(f"Candidate record rows: {len(record_map):,}")
    print(f"Excluded record rows: {len(excluded):,}")
    print(f"Unique AI payloads after de-duplication: {len(unique_rows):,}")
    print(f"AI task groups: {len(tasks):,}")


if __name__ == "__main__":
    main()
